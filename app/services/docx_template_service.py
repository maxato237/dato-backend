"""Génération d'un devis DATO à partir d'un template Word (.docx).

Deux familles de templates sont gérées (cf. ``render_quote_docx``) :

A. Template avec **tableau** dont la 1re ligne contient « Désignation » :
   on conserve l'en-tête du tableau et on remplace ses lignes de données.

B. Template **papèterie** (en-tête/pied de page Word seuls, corps vide) :
   on **construit tout le corps** du devis
   (titre, client, tableau, total en lettres, signatures) dans le corps,
   pendant que Word conserve l'en-tête et le pied de page sur chaque page.

Conversion finale en PDF via docx2pdf (Word COM, Windows/Mac) ou LibreOffice.

Dépendances :
    pip install python-docx docx2pdf
"""
from __future__ import annotations

import io
import os
import re
import subprocess
import tempfile
import unicodedata

from docx import Document  # type: ignore[import-untyped]
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT  # type: ignore[import-untyped]
from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore[import-untyped]
from docx.oxml import OxmlElement  # type: ignore[import-untyped]
from docx.oxml.ns import qn  # type: ignore[import-untyped]
from docx.shared import Cm, Emu, Pt, RGBColor  # type: ignore[import-untyped]
from docx.text.paragraph import Paragraph  # type: ignore[import-untyped]

from app.models.quote import Quote


# Mots-clés qui signalent la colonne « Désignation » dans la première ligne.
_HEADER_KEYWORDS = {
    'désignation', 'designation',
    'libellé', 'libelle',
    'description', 'article',
    'nature', 'objet',
}


# ---------------------------------------------------------------------------
# Localisation du tableau principal
# ---------------------------------------------------------------------------

def _find_items_table(doc: Document):
    """Retourne le tableau du devis dans le template, ou None."""
    for table in doc.tables:
        if not table.rows:
            continue
        first_row_text = ' '.join(
            c.text.strip().lower() for c in table.rows[0].cells
        )
        if any(kw in first_row_text for kw in _HEADER_KEYWORDS):
            return table

    # Fallback : tableau avec le plus grand nombre de colonnes (≥ 3).
    candidates = [t for t in doc.tables if len(t.columns) >= 3]
    if candidates:
        return max(candidates, key=lambda t: len(t.columns))
    return None


# ---------------------------------------------------------------------------
# Manipulation des lignes
# ---------------------------------------------------------------------------

def _delete_data_rows(table) -> None:
    """Supprime toutes les lignes du tableau sauf la première (en-tête)."""
    tbl_xml = table._tbl
    for row in list(table.rows)[1:]:
        tbl_xml.remove(row._tr)


def _set_cell_text(cell, text: str, bold: bool = False, align=None,
                   size: float | None = None, color: str | None = None) -> None:
    """Écrit `text` dans la cellule en réutilisant le style du premier run."""
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ''

    para = cell.paragraphs[0]
    if para.runs:
        run = para.runs[0]
    else:
        run = para.add_run()

    run.text = text
    if bold:
        run.bold = True
    elif run.bold:
        run.bold = False  # forcer non-gras si hérité
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if align is not None:
        para.alignment = align

    # Texte centré verticalement dans la hauteur de la cellule.
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _add_row(table, cells_data: list[str], bold: bool = False):
    """Ajoute une ligne avec les textes fournis."""
    row = table.add_row()
    for i, value in enumerate(cells_data):
        if i < len(row.cells):
            _set_cell_text(row.cells[i], value, bold=bold)
    return row


def _shade_cell(cell, fill_hex: str) -> None:
    """Applique une couleur de fond à une cellule (ex. 'E7EEF3')."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)


def _add_table_borders(table, color: str = 'C9D2DC', sz: str = '4') -> None:
    """Ajoute des bordures fines (grille complète) à un tableau."""
    tblPr = table._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), sz)
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        borders.append(el)
    tblPr.append(borders)


# ---------------------------------------------------------------------------
# Politique de pagination
# ---------------------------------------------------------------------------

def _row_cant_split(row) -> None:
    """Empêche la ligne de se couper en travers d'une page."""
    trPr = row._tr.get_or_add_trPr()
    trPr.append(OxmlElement('w:cantSplit'))


def _row_repeat_as_header(row) -> None:
    """La ligne se répète en haut de chaque page (en-tête de tableau)."""
    trPr = row._tr.get_or_add_trPr()
    trPr.append(OxmlElement('w:tblHeader'))


def _keep_with_next(cell) -> None:
    """Garde le contenu de la cellule collé à la ligne/au paragraphe suivant."""
    for para in cell.paragraphs:
        para.paragraph_format.keep_with_next = True


# ---------------------------------------------------------------------------
# Formatage
# ---------------------------------------------------------------------------

def _fmt(amount) -> str:
    """Formate un montant : 1 025 500 (séparateur milliers = espace)."""
    try:
        return f'{float(amount):,.0f}'.replace(',', ' ')  # espace fine
    except (TypeError, ValueError):
        return str(amount)


# ---------------------------------------------------------------------------
# API publique
# ---------------------------------------------------------------------------

def fill_quote_template(template_bytes: bytes, quote: Quote) -> bytes:
    """Remplit le template DOCX avec les données du devis.

    L'en-tête et le pied de page Word (styles natifs) sont préservés intacts.
    Seul le contenu du tableau principal est remplacé.

    Args:
        template_bytes: Contenu binaire du fichier .docx template.
        quote: Instance Quote SQLAlchemy chargée avec ses items et company.

    Returns:
        Contenu binaire du .docx rempli.

    Raises:
        ValueError: Si aucun tableau de devis n'est trouvé dans le template.
    """
    doc = Document(io.BytesIO(template_bytes))

    table = _find_items_table(doc)
    if table is None:
        raise ValueError(
            "Aucun tableau de devis trouvé dans le template. "
            "Le tableau doit contenir une colonne « Désignation » ou « Description »."
        )

    num_cols = len(table.columns)

    # ── 1. Vider les lignes de données ───────────────────────────────────────
    _delete_data_rows(table)

    # ── 2. Lignes d'articles ─────────────────────────────────────────────────
    for item in quote.items:
        desc = item.description or ''
        if item.unit:
            desc += f' ({item.unit})'

        cells = [''] * num_cols
        if num_cols >= 4:
            cells[0] = desc
            cells[1] = f'{float(item.quantity):g}'
            cells[2] = _fmt(item.unit_price)
            cells[3] = _fmt(item.total)
        elif num_cols == 3:
            cells[0] = desc
            cells[1] = f'{float(item.quantity):g}'
            cells[2] = _fmt(item.total)
        else:
            cells[0] = desc

        _add_row(table, cells)

    # ── 3. Sous-total ─────────────────────────────────────────────────────────
    if num_cols >= 2:
        sub = [''] * num_cols
        sub[0] = 'Sous-total'
        sub[-1] = _fmt(quote.subtotal)
        _add_row(table, sub, bold=True)

    # ── 4. TVA ────────────────────────────────────────────────────────────────
    if float(quote.tax_rate) > 0:
        tva = [''] * num_cols
        tva[0] = f'TVA ({float(quote.tax_rate):.0f} %)'
        tva[-1] = _fmt(quote.tax_amount)
        _add_row(table, tva, bold=True)

    # ── 5. Total général ─────────────────────────────────────────────────────
    tot = [''] * num_cols
    tot[0] = 'Total général'
    tot[-1] = _fmt(quote.total)
    _add_row(table, tot, bold=True)

    # ── 6. Sauvegarder ───────────────────────────────────────────────────────
    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


def _map_row_cells(kind: str, row: dict, num_cols: int) -> list[str]:
    """Mappe une ligne sémantique vers les cellules du tableau du template.

    Le template peut avoir 2, 3 ou 4+ colonnes ; l'app ne connaît pas cette
    structure, donc elle envoie des lignes sémantiques (« line » ou « span »)
    et c'est ici qu'on les place dans les bonnes colonnes.

    - kind == 'line' : article normal → Désignation / Qté / P.U / P.T
    - kind == 'span' : ligne de total/section → libellé (col 0) + montant (col -1)
    """
    cells = [''] * max(num_cols, 1)
    if kind == 'line':
        desig = str(row.get('designation', ''))
        qty = str(row.get('qty', ''))
        pu = str(row.get('pu', ''))
        pt = str(row.get('pt', ''))
        if num_cols >= 4:
            cells[0], cells[1], cells[2], cells[3] = desig, qty, pu, pt
        elif num_cols == 3:
            cells[0], cells[1], cells[2] = desig, qty, pt
        elif num_cols == 2:
            cells[0], cells[1] = desig, pt
        else:
            cells[0] = desig
    else:  # 'span' / 'section' : libellé + montant
        label = str(row.get('label', ''))
        amount = str(row.get('amount', ''))
        cells[0] = label
        if amount and num_cols >= 2:
            cells[-1] = amount
    return cells


def fill_template_rows(template_bytes: bytes, rows: list[dict]) -> bytes:
    """Remplit le tableau du template avec des lignes déjà formatées par l'app.

    L'en-tête et le pied de page Word sont préservés intacts ; seules les
    lignes de données du tableau principal sont remplacées.

    Args:
        template_bytes: contenu binaire du .docx template.
        rows: liste de lignes sémantiques. Chaque ligne est un dict :
              {'kind': 'line', 'designation', 'qty', 'pu', 'pt'} ou
              {'kind': 'span', 'label', 'amount', 'bold': bool}.

    Returns:
        Contenu binaire du .docx rempli.

    Raises:
        ValueError: si aucun tableau de devis n'est trouvé dans le template.
    """
    doc = Document(io.BytesIO(template_bytes))

    table = _find_items_table(doc)
    if table is None:
        raise ValueError(
            "Aucun tableau de devis trouvé dans le template. "
            "Le tableau doit contenir une colonne « Désignation » ou « Description »."
        )

    _fill_existing_table(table, rows)

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


# ---------------------------------------------------------------------------
# Reprise du design du tableau du modèle (approche A)
# ---------------------------------------------------------------------------
#
# But : conserver la mise en forme du tableau dessinée dans le modèle Word
# (fonds des bandes de sous-totaux, fond du total général, police), au lieu de
# ré-injecter des lignes neutres.
#
# Deux sources de style, par ordre de priorité :
#
# 1. Marqueurs explicites — une ligne-gabarit dont la 1re cellule commence par
#    un marqueur, mise en forme comme voulu :
#       @ligne          → style des lignes d'articles
#       @total          → style des sous-totaux (sections, rubriques)
#       @total-general  → style du total général (mis en avant)
#    Marqueurs insensibles à la casse et aux accents (« @TOTAL GÉNÉRAL » OK).
#
# 2. Détection automatique (si aucun marqueur) — on lit les lignes déjà
#    stylées du modèle : les lignes teintées deviennent le style des bandes
#    (sous-total / total général selon le libellé), la 1re ligne non teintée
#    le style des lignes d'articles. Permet de respecter un modèle déjà conçu
#    sans rien y ajouter. Si le modèle n'a aucune ligne stylée, rendu neutre.

_MARKER_KIND = {
    'ligne': 'line',
    'line': 'line',
    'total': 'span',
    'soustotal': 'span',
    'totalgeneral': 'strong',
}


class _RowStyle:
    """Mise en forme capturée d'une ligne-gabarit du modèle."""

    def __init__(self, fill=None, bold=None, color=None, size=None):
        self.fill = fill      # fond hex (ex. 'EEF3EE') ou None
        self.bold = bold      # bool ou None
        self.color = color    # couleur de police hex ou None
        self.size = size      # taille en points ou None


def _norm_marker(text: str) -> str:
    """Minuscules, sans accents ni caractères non alphanumériques."""
    text = unicodedata.normalize('NFKD', text)
    text = ''.join(c for c in text if not unicodedata.combining(c))
    return re.sub(r'[^a-z0-9]', '', text.lower())


def _marker_kind(cell_text: str):
    """Type de ligne ('line'/'span'/'strong') si la cellule est un marqueur
    « @… », sinon None."""
    t = cell_text.strip()
    if not t.startswith('@'):
        return None
    return _MARKER_KIND.get(_norm_marker(t))


def _cell_fill(cell):
    """Couleur de fond (hex) d'une cellule, ou None."""
    tcPr = cell._tc.tcPr
    if tcPr is None:
        return None
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        return None
    fill = shd.get(qn('w:fill'))
    return fill if fill and fill != 'auto' else None


def _first_run_style(cell):
    """(bold, color_hex|None, size_pt|None) du 1er run de la cellule."""
    for para in cell.paragraphs:
        for run in para.runs:
            color = None
            try:
                if run.font.color is not None and run.font.color.rgb is not None:
                    color = str(run.font.color.rgb)
            except (AttributeError, ValueError):
                color = None
            size = run.font.size.pt if run.font.size is not None else None
            return bool(run.bold), color, size
    return None, None, None


def _extract_row_style(row) -> _RowStyle:
    """Capture fond + police d'une ligne-gabarit."""
    fill = None
    for cell in row.cells:
        fill = _cell_fill(cell)
        if fill:
            break
    bold, color, size = (
        _first_run_style(row.cells[0]) if row.cells else (None, None, None))
    return _RowStyle(fill=fill, bold=bold, color=color, size=size)


def _capture_template_styles(table) -> dict:
    """Styles à reprendre du modèle → {kind: _RowStyle}, kind ∈ line/span/strong.

    Priorité aux marqueurs explicites (« @… ») ; à défaut, détection auto des
    lignes déjà stylées. À appeler avant ``_delete_data_rows`` (les lignes lues
    sont supprimées ensuite avec les autres données)."""
    styles: dict = {}
    for row in list(table.rows)[1:]:
        if not row.cells:
            continue
        kind = _marker_kind(row.cells[0].text)
        if kind and kind not in styles:
            styles[kind] = _extract_row_style(row)
    if styles:
        return styles
    return _infer_template_styles(table)


def _infer_template_styles(table) -> dict:
    """Déduit les styles du modèle depuis ses lignes déjà mises en forme.

    - ligne teintée dont le libellé contient « général » → total général ;
    - autre ligne teintée → bande de sous-total ;
    - 1re ligne non teintée avec du texte → ligne d'article.
    Repli : si aucune ligne « général », la dernière ligne teintée fait office
    de total général (souvent le cas dans les modèles)."""
    styles: dict = {}
    last_shaded = None
    for row in list(table.rows)[1:]:
        if not row.cells:
            continue
        fill = None
        for cell in row.cells:
            fill = _cell_fill(cell)
            if fill:
                break
        if fill:
            last_shaded = row
            label = _norm_marker(row.cells[0].text)
            if 'general' in label:
                styles.setdefault('strong', _extract_row_style(row))
            else:
                styles.setdefault('span', _extract_row_style(row))
        elif row.cells[0].text.strip():
            styles.setdefault('line', _extract_row_style(row))
    if 'strong' not in styles and last_shaded is not None:
        styles['strong'] = _extract_row_style(last_shaded)
    return styles


def _apply_row_style(row, style) -> None:
    """Applique le fond + la police d'un style capturé à une ligne générée."""
    if style is None:
        return
    seen = set()
    for cell in row.cells:
        # Après une fusion, ``row.cells`` répète la cellule fusionnée : on ne la
        # traite qu'une fois (sinon w:shd dupliqué).
        tc = cell._tc
        if tc in seen:
            continue
        seen.add(tc)
        if style.fill:
            _shade_cell(cell, style.fill)
        for para in cell.paragraphs:
            for run in para.runs:
                if style.bold is not None:
                    run.bold = style.bold
                if style.color:
                    run.font.color.rgb = RGBColor.from_string(style.color)
                if style.size:
                    run.font.size = Pt(style.size)


def _fill_existing_table(table, rows: list[dict]) -> None:
    """Remplace les lignes de données du tableau du modèle par `rows`.

    Le modèle ne donne que le **style** du tableau (fonds des bandes, police,
    en-tête) ; le **contenu** vient de l'app. On reprend donc le design (via
    lignes-gabarits « @… » ou détection auto), on garantit une grille de
    bordures, et on aligne les cellules (qté centrée, montants à droite,
    centrage vertical)."""
    num_cols = len(table.columns)
    styles = _capture_template_styles(table)
    _delete_data_rows(table)
    _ensure_grid_borders(table)
    # En-tête conservée : répétée sur chaque page + non sécable.
    if table.rows:
        _row_cant_split(table.rows[0])
        _row_repeat_as_header(table.rows[0])
    for row in rows:
        kind = str(row.get('kind', 'line'))
        bold = bool(row.get('bold', False))
        strong = bool(row.get('strong', False))
        if kind == 'span' and num_cols >= 2:
            style = (styles.get('strong') or styles.get('span')) if strong \
                else styles.get('span')
            new_row = _add_span_row(
                table, num_cols, str(row.get('label', '')),
                str(row.get('amount', '')), bold=bold)
        else:
            style = styles.get('line')
            new_row = _add_line_row(table, num_cols, row)
        _row_cant_split(new_row)
        _apply_row_style(new_row, style)


def _ensure_grid_borders(table) -> None:
    """Force une grille de bordures fines (toutes les lignes encadrées)."""
    tblPr = table._tbl.tblPr
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    _add_table_borders(table, color='BFBFBF', sz='4')


def _add_line_row(table, num_cols: int, row: dict):
    """Ajoute une ligne d'article : seule la désignation (texte) reste à
    gauche ; tous les nombres (Qté, P.U, P.T) sont alignés à droite."""
    new_row = table.add_row()
    cells = new_row.cells
    desig = str(row.get('designation', ''))
    qty = str(row.get('qty', ''))
    pu = str(row.get('pu', ''))
    pt = str(row.get('pt', ''))
    R = WD_ALIGN_PARAGRAPH.RIGHT
    if num_cols >= 4:
        _set_cell_text(cells[0], desig)
        _set_cell_text(cells[1], qty, align=R)
        _set_cell_text(cells[2], pu, align=R)
        _set_cell_text(cells[3], pt, align=R)
    elif num_cols == 3:
        _set_cell_text(cells[0], desig)
        _set_cell_text(cells[1], qty, align=R)
        _set_cell_text(cells[2], pt, align=R)
    elif num_cols == 2:
        _set_cell_text(cells[0], desig)
        _set_cell_text(cells[1], pt, align=R)
    else:
        _set_cell_text(cells[0], desig)
    return new_row


def _add_span_row(table, num_cols: int, label: str, amount: str,
                  bold: bool = False):
    """Ajoute une bande (sous-total / total) : libellé fusionné + montant à
    droite, comme les lignes de total dessinées dans le modèle."""
    new_row = table.add_row()
    cells = new_row.cells
    label_cell = cells[0]
    for i in range(1, num_cols - 1):
        label_cell = label_cell.merge(cells[i])
    _set_cell_text(label_cell, label, bold=bold,
                   align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell_text(new_row.cells[-1], amount, bold=bold,
                   align=WD_ALIGN_PARAGRAPH.RIGHT)
    return new_row


# ---------------------------------------------------------------------------
# Construction complète du corps (template « papèterie », corps vide)
# ---------------------------------------------------------------------------

_BLUE = '1B4965'
_HEAD_BG = 'E7EEF3'
_SPAN_BG = 'EEF3EE'


def _build_items_table(doc, rows: list[dict]) -> None:
    """Construit le tableau Désignation / Qté / P.U / P.T depuis zéro."""
    table = doc.add_table(rows=1, cols=4)
    _add_table_borders(table)

    headers = ['Désignation', 'Qté', 'P.U', 'P.T']
    aligns = [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER,
              WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER]
    for i, head in enumerate(headers):
        cell = table.rows[0].cells[i]
        _set_cell_text(cell, head, bold=True, align=aligns[i], size=10)
        _shade_cell(cell, _HEAD_BG)
    # En-tête répétée en haut de chaque page + non sécable.
    _row_cant_split(table.rows[0])
    _row_repeat_as_header(table.rows[0])

    for row in rows:
        kind = str(row.get('kind', 'line'))
        bold = bool(row.get('bold', False))
        strong = bool(row.get('strong', False))  # ligne mise en avant (total général)
        new_row = table.add_row()
        cells = new_row.cells
        _row_cant_split(new_row)  # jamais couper une ligne en travers d'une page
        if kind == 'line':
            _set_cell_text(cells[0], str(row.get('designation', '')), size=10)
            _set_cell_text(cells[1], str(row.get('qty', '')),
                           align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
            _set_cell_text(cells[2], str(row.get('pu', '')),
                           align=WD_ALIGN_PARAGRAPH.RIGHT, size=10)
            _set_cell_text(cells[3], str(row.get('pt', '')),
                           align=WD_ALIGN_PARAGRAPH.RIGHT, size=10)
        else:  # span : libellé fusionné sur 3 colonnes + montant
            label_cell = cells[0].merge(cells[1]).merge(cells[2])
            if strong:
                # Total général : fond bleu foncé, texte blanc, plus grand.
                _set_cell_text(label_cell, str(row.get('label', '')),
                               bold=True, size=11, color='FFFFFF')
                _set_cell_text(cells[3], str(row.get('amount', '')),
                               bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT,
                               size=11, color='FFFFFF')
                _shade_cell(label_cell, _BLUE)
                _shade_cell(cells[3], _BLUE)
            else:
                _set_cell_text(label_cell, str(row.get('label', '')),
                               bold=True, size=10)
                _set_cell_text(cells[3], str(row.get('amount', '')),
                               bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT, size=10)
                if bold:
                    _shade_cell(label_cell, _SPAN_BG)
                    _shade_cell(cells[3], _SPAN_BG)
                # En-tête de section (span sans montant) : reste collée à sa
                # 1re ligne (pas de titre orphelin en bas de page).
                if not str(row.get('amount', '')).strip():
                    _keep_with_next(label_cell)


def _build_quote_body(doc, data: dict) -> None:
    """Construit tout le corps du devis dans un template papèterie (corps vide).

    Reproduit la mise en page du rendu par défaut de l'app : ligne de date,
    titre, « DOIT : client », tableau, total en lettres, NB, signatures.
    """
    # Ligne ville + date (alignée à droite).
    city_date = str(data.get('cityDate', '')).strip()
    if city_date:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(18)
        p.add_run(city_date)

    # Titre centré, gras, souligné.
    title = str(data.get('title', '')).strip()
    if title:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(16)
        run = p.add_run(title)
        run.bold = True
        run.underline = True
        run.font.size = Pt(13)

    # « DOIT : {client} ».
    client = str(data.get('client', '')).strip()
    if client:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(12)
        r1 = p.add_run('DOIT : ')
        r1.bold = True
        r1.underline = True
        r2 = p.add_run(client)
        r2.bold = True

    # Tableau des articles.
    _build_items_table(doc, data.get('rows', []))

    # Bloc de conclusion (Arrêté… + NB + signatures) soudé : il bascule en
    # entier sur la page suivante plutôt que de laisser les signatures seules.
    # « Arrêté le présent devis à la somme de … ».
    words = str(data.get('amountInWords', '')).strip().upper()
    if words:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        r1 = p.add_run('Arrêté le présent devis à la somme de ')
        r1.italic = True
        r2 = p.add_run(words + '.')
        r2.italic = True
        r2.bold = True

    # NB / note.
    note = str(data.get('note', '')).strip()
    if note:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        r1 = p.add_run('NB : ')
        r1.bold = True
        p.add_run(note)

    # Signatures (une colonne par signataire).
    sigs = [str(s) for s in (data.get('signatures') or []) if str(s).strip()]
    if sigs:
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(18)
        spacer.paragraph_format.keep_with_next = True
        st = doc.add_table(rows=1, cols=len(sigs))
        _row_cant_split(st.rows[0])  # les colonnes de signature restent ensemble
        for i, label in enumerate(sigs):
            cell = st.rows[0].cells[i]
            _set_cell_text(cell, label, bold=True,
                           align=WD_ALIGN_PARAGRAPH.CENTER)
            # Un seul espace de signature (gap modéré entre le libellé et
            # le mot « Signature »).
            gap = cell.add_paragraph()
            gap.paragraph_format.space_before = Pt(20)
            sp = cell.add_paragraph('Signature')
            sp.alignment = WD_ALIGN_PARAGRAPH.CENTER


# ---------------------------------------------------------------------------
# Remplacement des placeholders du modèle (mode « modèle avec tableau »)
# ---------------------------------------------------------------------------
#
# Un modèle « avec tableau » ne voit que son tableau rempli : le titre, le
# client et le montant en lettres restent le texte statique du modèle. Pour les
# rendre dynamiques sans imposer de mise en page, on remplace des placeholders
# « {{…}} » par les valeurs réelles du devis, où qu'ils soient dans le corps
# (paragraphes, cellules, en-tête/pied de page Word) :
#
#   {{titre}} {{client}} {{montant_en_lettres}} {{date}} {{numero}} {{note}}
#
# Insensibles à la casse/accents/underscores. Placeholder inconnu : laissé tel
# quel (on ne casse rien).

_PLACEHOLDER_KEYS = {
    'titre': 'title', 'title': 'title', 'objet': 'title',
    'client': 'client', 'doit': 'client',
    'montantenlettres': 'amountInWords', 'lettres': 'amountInWords',
    'date': 'cityDate', 'villedate': 'cityDate',
    'numero': 'number', 'number': 'number',
    'note': 'note', 'nb': 'note',
}

_PLACEHOLDER_RE = re.compile(r'\{\{([^{}]+)\}\}')


def _replace_in_paragraph(paragraph, resolve) -> None:
    """Remplace les « {{…}} » d'un paragraphe via ``resolve(match)``."""
    runs = paragraph.runs
    if not runs:
        return
    # 1. Placeholder contenu dans un seul run : remplacement local, le
    #    formatage du run est préservé.
    for run in runs:
        if '{{' in run.text and '}}' in run.text:
            run.text = _PLACEHOLDER_RE.sub(resolve, run.text)
    # 2. Placeholder éclaté sur plusieurs runs : on recompose le paragraphe
    #    dans le 1er run (Word fragmente souvent un mot en plusieurs runs).
    full = ''.join(r.text for r in runs)
    if _PLACEHOLDER_RE.search(full):
        runs[0].text = _PLACEHOLDER_RE.sub(resolve, full)
        for r in runs[1:]:
            r.text = ''


def _replace_placeholders(doc, data: dict) -> None:
    """Remplace les placeholders « {{…}} » du corps par les valeurs du devis."""
    def resolve(match):
        key = _PLACEHOLDER_KEYS.get(_norm_marker(match.group(1)))
        if key is None:
            return match.group(0)  # placeholder inconnu : laissé intact
        val = '' if data.get(key) is None else str(data.get(key)).strip()
        # La note porte son étiquette « NB : » uniquement si elle est remplie
        # (placeholder vide ⇒ pas d'étiquette orpheline).
        if key == 'note':
            return f'NB : {val}' if val else ''
        return val

    paragraphs = list(doc.paragraphs)
    for table in doc.tables:
        for trow in table.rows:
            for cell in trow.cells:
                paragraphs.extend(cell.paragraphs)
    for section in doc.sections:
        paragraphs.extend(section.header.paragraphs)
        paragraphs.extend(section.footer.paragraphs)
    for para in paragraphs:
        _replace_in_paragraph(para, resolve)


# ---------------------------------------------------------------------------
# Injection du contenu de l'app dans le corps du modèle (sans placeholders)
# ---------------------------------------------------------------------------
#
# Principe : le modèle ne donne QUE le style. Le corps du modèle est un exemple
# de devis ; on remplace son texte par celui du devis de l'app **en place**
# (donc en gardant la mise en forme de chaque paragraphe), en reconnaissant
# chaque ligne par son motif :
#   - « DOIT … »                  → DOIT : {client}
#   - « Arrêté … »                → Arrêté … à la somme de {montant_en_lettres}.
#   - « …, le 22 Juin 2026 »      → {cityDate} (ville + date du devis)
#   - titre (centré + gras, « DEVIS ») → {title}
# La note (NB) est insérée après la ligne « Arrêté… » si elle est renseignée.

_DATE_RE = re.compile(r'\ble\s+\d{1,2}\s+\w+\s+\d{4}', re.IGNORECASE)


def _has_placeholders(doc) -> bool:
    return any('{{' in p.text for p in doc.paragraphs)


def _set_para_text(paragraph, text: str) -> None:
    """Réduit le paragraphe à son 1er run (format conservé) avec `text`."""
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for r in paragraph.runs[1:]:
        r.text = ''


def _looks_like_title(paragraph) -> bool:
    txt = paragraph.text.strip()
    if len(txt) < 15 or 'DEVIS' not in txt.upper():
        return False
    centered = paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER
    bold = any(r.bold for r in paragraph.runs)
    return centered and bold


def _apply_body_content(doc, data: dict) -> None:
    """Injecte le contenu du devis dans le corps du modèle, en place."""
    title = str(data.get('title', '')).strip()
    client = str(data.get('client', '')).strip()
    city_date = str(data.get('cityDate', '')).strip()
    words = str(data.get('amountInWords', '')).strip().upper()
    note = str(data.get('note', '')).strip()

    arrete_para = None
    done = set()
    strays = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if not t:
            continue
        norm = t.replace('\xa0', ' ')
        if 'doit' not in done and norm.upper().startswith('DOIT'):
            # Garde « DOIT : » du modèle, ajoute le client en gras.
            _set_para_text(p, 'DOIT : ')
            run = p.add_run(client)
            run.bold = True
            done.add('doit')
        elif 'arrete' not in done and norm.lower().startswith('arrêté'):
            runs = p.runs
            runs[0].text = 'Arrêté le présent devis à la somme de '
            if len(runs) > 1:
                runs[1].text = words + '.'
                for r in runs[2:]:
                    r.text = ''
            else:
                extra = p.add_run(words + '.')
                extra.bold = True
                extra.italic = True
            arrete_para = p
            done.add('arrete')
        elif 'date' not in done and city_date and _DATE_RE.search(norm):
            _set_para_text(p, city_date)
            done.add('date')
        elif 'title' not in done and _looks_like_title(p):
            _set_para_text(p, title)
            done.add('title')
        elif (p.alignment == WD_ALIGN_PARAGRAPH.CENTER
              and 'DEVIS' in norm.upper()
              and not any(r.bold for r in p.runs)):
            # Fragment de titre résiduel du modèle (non gras) : contenu parasite.
            strays.append(p)

    for p in strays:
        p._p.getparent().remove(p._p)

    # NB inséré après la ligne « Arrêté… » (uniquement si renseigné).
    if note and arrete_para is not None:
        new_p = OxmlElement('w:p')
        arrete_para._p.addnext(new_p)
        np = Paragraph(new_p, arrete_para._parent)
        run = np.add_run(f'NB : {note}')
        run.bold = False


def _ensure_header_clearance(doc, gap_cm: float = 0.6) -> None:
    """Évite que le corps (et l'en-tête de tableau répété sur les pages
    suivantes) ne se colle au bandeau d'en-tête Word.

    Si la marge haute est plus petite que la hauteur du bandeau d'en-tête,
    on l'augmente à « hauteur du bandeau + marge » pour ménager un interligne.
    On n'augmente jamais à la baisse (un modèle déjà bien réglé est préservé)."""
    wp = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
    for section in doc.sections:
        heights = [
            int(e.get('cy'))
            for e in section.header._element.iter(f'{{{wp}}}extent')
            if e.get('cy')
        ]
        if not heights:
            continue
        needed = max(heights) + Cm(gap_cm)
        if section.top_margin is None or section.top_margin < needed:
            section.top_margin = Emu(int(needed))


def render_quote_docx(template_bytes: bytes, data: dict) -> bytes:
    """Produit le .docx final du devis à partir du template et des données app.

    - Modèle **avec tableau** : le modèle ne fournit que le style. On reprend le
      design du tableau, on le remplit avec les lignes de l'app, et on injecte
      le contenu (titre, date, client, montant en lettres, NB) dans le corps —
      par placeholders « {{…}} » si le modèle en contient, sinon par
      reconnaissance des lignes du corps.
    - Modèle **papèterie** (sans tableau) : on construit tout le corps.

    `data` : {'rows', 'cityDate', 'title', 'client', 'amountInWords',
              'note', 'signatures', 'number'}.
    """
    doc = Document(io.BytesIO(template_bytes))

    table = _find_items_table(doc)
    if table is not None:
        _fill_existing_table(table, data.get('rows', []))
        if _has_placeholders(doc):
            _replace_placeholders(doc, data)
        else:
            _apply_body_content(doc, data)
    else:
        _build_quote_body(doc, data)

    # Marge haute suffisante pour ne pas coller le contenu au bandeau d'en-tête
    # (visible surtout sur les pages 2+ où l'en-tête de tableau est répété).
    _ensure_header_clearance(doc)

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


def docx_to_pdf(docx_bytes: bytes) -> bytes:
    """Convertit un .docx en PDF.

    Essaie dans l'ordre :
      1. ``docx2pdf`` (Windows/Mac avec Microsoft Word installé)
      2. LibreOffice headless (Linux ou si Word absent)

    Returns:
        Contenu binaire du PDF généré.

    Raises:
        RuntimeError: Si aucun convertisseur n'est disponible.
    """
    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = os.path.join(tmpdir, 'quote.docx')
        pdf_path  = os.path.join(tmpdir, 'quote.pdf')

        with open(docx_path, 'wb') as fh:
            fh.write(docx_bytes)

        # ── Tentative 1 : docx2pdf (Word COM / AppleScript) ──────────────────
        # Word COM doit être initialisé dans le thread courant (Flask sert
        # chaque requête dans un thread worker → sinon « CoInitialize has not
        # been called »).
        _com_ready = False
        try:
            import pythoncom  # type: ignore[import-untyped]
            pythoncom.CoInitialize()
            _com_ready = True
        except Exception:
            pass
        try:
            from docx2pdf import convert  # type: ignore[import-untyped]
            convert(docx_path, pdf_path)
            if os.path.exists(pdf_path):
                with open(pdf_path, 'rb') as fh:
                    return fh.read()
        except Exception:
            pass
        finally:
            if _com_ready:
                try:
                    pythoncom.CoUninitialize()
                except Exception:
                    pass

        # ── Tentative 2 : LibreOffice headless ───────────────────────────────
        _SOFFICE_CANDIDATES = [
            'soffice', 'libreoffice',
            '/usr/bin/soffice', '/usr/bin/libreoffice',
            '/usr/local/bin/soffice',
        ]
        for soffice in _SOFFICE_CANDIDATES:
            try:
                subprocess.run(
                    [soffice, '--headless', '--convert-to', 'pdf',
                     '--outdir', tmpdir, docx_path],
                    capture_output=True,
                    timeout=60,
                    check=True,
                )
                if os.path.exists(pdf_path):
                    with open(pdf_path, 'rb') as fh:
                        return fh.read()
            except (FileNotFoundError, subprocess.CalledProcessError,
                    subprocess.TimeoutExpired):
                continue

        raise RuntimeError(
            "Impossible de convertir le DOCX en PDF. "
            "Installez Microsoft Word (Windows/Mac) ou LibreOffice, "
            "puis relancez le backend."
        )
